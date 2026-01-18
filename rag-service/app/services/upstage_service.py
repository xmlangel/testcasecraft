"""Document Analysis Service with Multiple Parser Support"""
from typing import List, Dict, Any, Optional
import logging
import httpx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from ..core.config import settings

logger = logging.getLogger(__name__)


class UpstageService:
    """Service for Document Analysis with Multiple Parser Options"""

    def __init__(self, parser: str = None, api_key: str = None):
        """
        Initialize document parser based on configuration or parameters

        Args:
            parser: Parser to use (upstage, pymupdf, pymupdf4llm, pypdf, auto).
                   If None, uses settings.DOCUMENT_PARSER
            api_key: Upstage API key. If None, uses settings.UPSTAGE_API_KEY
        """
        self.parser_mode = parser or settings.DOCUMENT_PARSER
        self.api_key = api_key or settings.UPSTAGE_API_KEY
        self.api_url = "https://api.upstage.ai/v1/document-ai/layout-analysis"

        # Determine actual parser to use
        if self.parser_mode == "auto":
            # Auto-select priority: upstage > pymupdf4llm > pymupdf > pypdf
            if self.api_key:
                self.active_parser = "upstage"
            else:
                # Try to import libraries to determine best local parser
                try:
                    import pymupdf4llm
                    self.active_parser = "pymupdf4llm"
                except ImportError:
                    try:
                        import fitz  # PyMuPDF
                        self.active_parser = "pymupdf"
                    except ImportError:
                        self.active_parser = "pypdf"  # Fallback to pypdf
        else:
            self.active_parser = self.parser_mode

        # Text splitter configuration
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

        logger.info(f"UpstageService initialized (mode: {self.parser_mode}, active: {self.active_parser})")

    async def analyze_document(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """
        Analyze document using configured parser

        Args:
            file_content: Document file content as bytes
            file_name: Original file name

        Returns:
            Analysis result with structured content
        """
        try:
            # Determine file type
            file_ext = file_name.lower().split('.')[-1]

            if file_ext not in ['pdf', 'docx', 'doc', 'txt']:
                raise ValueError(f"Unsupported file type: {file_ext}")

            # Route to appropriate parser
            if self.active_parser == "upstage":
                return await self._parse_with_upstage(file_content, file_name, file_ext)
            elif self.active_parser == "pymupdf4llm":
                return await self._parse_with_pymupdf4llm(file_content, file_name, file_ext)
            elif self.active_parser == "pymupdf":
                return await self._parse_with_pymupdf(file_content, file_name, file_ext)
            else:  # pypdf
                return await self._parse_with_pypdf(file_content, file_name, file_ext)

        except Exception as e:
            logger.error(f"Error analyzing document {file_name}: {str(e)}")
            raise Exception(f"Document analysis failed: {str(e)}")

    async def _parse_with_upstage(self, file_content: bytes, file_name: str, file_ext: str) -> Dict[str, Any]:
        """Parse document using Upstage Layout Analysis API"""
        try:
            logger.info(f"Parsing with Upstage API: {file_name}")

            import io

            async with httpx.AsyncClient() as client:
                # Upstage API expects 'document' field as file-like object
                # Use BytesIO to create file-like object from bytes
                file_obj = io.BytesIO(file_content)

                # Simple file upload format (httpx will set correct Content-Type)
                files = {"document": (file_name, file_obj)}
                headers = {"Authorization": f"Bearer {self.api_key}"}

                # Optional: Add OCR parameter
                data = {"ocr": "true"}

                response = await client.post(
                    self.api_url,
                    files=files,
                    headers=headers,
                    data=data,
                    timeout=60.0
                )

                response.raise_for_status()
                data = response.json()

            # Extract structured information
            result = {
                "text": data.get("text", ""),
                "html": data.get("html", ""),
                "markdown": data.get("markdown", ""),
                "metadata": {
                    "file_name": file_name,
                    "file_type": file_ext,
                    "pages": data.get("pages", 0),
                    "elements": data.get("elements", []),
                    "parser": "upstage"
                }
            }

            logger.info(f"Upstage parsing completed: {file_name}")
            return result

        except httpx.HTTPError as e:
            logger.error(f"Upstage API error: {str(e)}")
            raise Exception(f"Upstage API request failed: {str(e)}")

    async def _parse_with_pymupdf4llm(self, file_content: bytes, file_name: str, file_ext: str) -> Dict[str, Any]:
        """Parse document using PyMuPDF4LLM (LLM-optimized markdown)"""
        try:
            logger.info(f"Parsing with PyMuPDF4LLM: {file_name}")

            import io
            import pymupdf4llm

            text = ""
            markdown = ""
            pages = 0

            if file_ext == 'pdf':
                # Use PyMuPDF4LLM for LLM-optimized markdown extraction
                # PyMuPDF4LLM requires a file path, so save to temp file
                import tempfile
                import os

                # Create temporary file and write content
                temp_fd, temp_path = tempfile.mkstemp(suffix='.pdf')
                try:
                    # Write content to temp file and close it
                    os.write(temp_fd, file_content)
                    os.close(temp_fd)

                    # Now use the closed temp file with PyMuPDF4LLM
                    markdown_result = pymupdf4llm.to_markdown(temp_path)

                    if isinstance(markdown_result, str):
                        markdown = markdown_result
                        text = markdown_result  # Use markdown as text too
                    elif isinstance(markdown_result, dict):
                        markdown = markdown_result.get("markdown", "")
                        text = markdown_result.get("text", markdown)
                        pages = markdown_result.get("pages", 0)

                    # Get page count if not already set
                    if pages == 0:
                        import fitz
                        doc = fitz.open(temp_path)
                        pages = len(doc)
                        doc.close()

                    logger.info(f"Extracted {len(text)} characters from {pages} pages (PyMuPDF4LLM)")
                finally:
                    # Clean up temporary file
                    if os.path.exists(temp_path):
                        os.unlink(temp_path)

            elif file_ext in ['doc', 'docx']:
                # Fallback to python-docx for DOCX
                from docx import Document
                docx_file = io.BytesIO(file_content)
                doc = Document(docx_file)
                pages = 1
                text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
                text = "\n\n".join(text_parts)
                markdown = text
                logger.info(f"Extracted {len(text)} characters from DOCX (python-docx)")

            elif file_ext == 'txt':
                text = file_content.decode('utf-8', errors='ignore')
                markdown = text
                pages = 1

            result = {
                "text": text,
                "html": "",
                "markdown": markdown,
                "metadata": {
                    "file_name": file_name,
                    "file_type": file_ext,
                    "pages": pages,
                    "elements": [],
                    "parser": "pymupdf4llm"
                }
            }

            logger.info(f"PyMuPDF4LLM parsing completed: {file_name}")
            return result

        except Exception as e:
            logger.error(f"PyMuPDF4LLM parsing error: {str(e)}")
            raise Exception(f"PyMuPDF4LLM parsing failed: {str(e)}")

    async def _parse_with_pymupdf(self, file_content: bytes, file_name: str, file_ext: str) -> Dict[str, Any]:
        """Parse document using PyMuPDF (fitz) - fast and feature-rich"""
        try:
            logger.info(f"Parsing with PyMuPDF: {file_name}")

            import io
            import fitz  # PyMuPDF

            text = ""
            pages = 0

            if file_ext == 'pdf':
                # Use PyMuPDF for PDF parsing
                pdf_file = io.BytesIO(file_content)
                doc = fitz.open(stream=pdf_file, filetype="pdf")
                pages = len(doc)

                # Extract text from all pages
                text_parts = []
                for page in doc:
                    page_text = page.get_text()
                    if page_text:
                        text_parts.append(page_text)

                text = "\n\n".join(text_parts)
                doc.close()
                logger.info(f"Extracted {len(text)} characters from {pages} pages (PyMuPDF)")

            elif file_ext in ['doc', 'docx']:
                # Fallback to python-docx for DOCX
                from docx import Document
                docx_file = io.BytesIO(file_content)
                doc = Document(docx_file)
                pages = 1
                text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
                text = "\n\n".join(text_parts)
                logger.info(f"Extracted {len(text)} characters from DOCX (python-docx)")

            elif file_ext == 'txt':
                text = file_content.decode('utf-8', errors='ignore')
                pages = 1

            result = {
                "text": text,
                "html": "",
                "markdown": text,
                "metadata": {
                    "file_name": file_name,
                    "file_type": file_ext,
                    "pages": pages,
                    "elements": [],
                    "parser": "pymupdf"
                }
            }

            logger.info(f"PyMuPDF parsing completed: {file_name}")
            return result

        except Exception as e:
            logger.error(f"PyMuPDF parsing error: {str(e)}")
            raise Exception(f"PyMuPDF parsing failed: {str(e)}")

    async def _parse_with_pypdf(self, file_content: bytes, file_name: str, file_ext: str) -> Dict[str, Any]:
        """Parse document using pypdf and python-docx (basic parsers)"""
        try:
            logger.info(f"Parsing with pypdf: {file_name}")

            import io

            text = ""
            pages = 0

            if file_ext == 'pdf':
                # Use PyPDF2 for PDF parsing
                from pypdf import PdfReader
                pdf_file = io.BytesIO(file_content)
                reader = PdfReader(pdf_file)
                pages = len(reader.pages)

                # Extract text from all pages
                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                text = "\n\n".join(text_parts)
                logger.info(f"Extracted {len(text)} characters from {pages} pages (pypdf)")

            elif file_ext in ['doc', 'docx']:
                # Use python-docx for DOCX parsing
                from docx import Document
                docx_file = io.BytesIO(file_content)
                doc = Document(docx_file)
                pages = 1
                text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
                text = "\n\n".join(text_parts)
                logger.info(f"Extracted {len(text)} characters from DOCX (python-docx)")

            elif file_ext == 'txt':
                text = file_content.decode('utf-8', errors='ignore')
                pages = 1

            result = {
                "text": text,
                "html": "",
                "markdown": text,
                "metadata": {
                    "file_name": file_name,
                    "file_type": file_ext,
                    "pages": pages,
                    "elements": [],
                    "parser": "pypdf"
                }
            }

            logger.info(f"pypdf parsing completed: {file_name}")
            return result

        except Exception as e:
            logger.error(f"pypdf parsing error: {str(e)}")
            raise Exception(f"pypdf parsing failed: {str(e)}")

    def create_chunks(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """
        Split text into semantic chunks

        Args:
            text: Full document text
            metadata: Optional metadata to attach to chunks

        Returns:
            List of chunk dictionaries with text and metadata
        """
        try:
            if not text:
                logger.warning("Empty text provided for chunking")
                return []

            # Split text into chunks
            chunks = self.text_splitter.split_text(text)

            # Create chunk objects with metadata
            chunk_objects = []
            for idx, chunk_text in enumerate(chunks):
                chunk_metadata = metadata.copy() if metadata else {}
                chunk_metadata.update({
                    "chunk_index": idx,
                    "chunk_size": len(chunk_text)
                })

                chunk_objects.append({
                    "index": idx,
                    "text": chunk_text,
                    "metadata": chunk_metadata
                })

            logger.info(f"Created {len(chunk_objects)} chunks from text (length: {len(text)})")
            return chunk_objects

        except Exception as e:
            logger.error(f"Error creating chunks: {str(e)}")
            raise Exception(f"Chunk creation failed: {str(e)}")

    async def analyze_and_chunk(self, file_content: bytes, file_name: str) -> Dict[str, Any]:
        """
        Analyze document and create chunks in one operation

        Args:
            file_content: Document file content
            file_name: Original file name

        Returns:
            Dictionary with analysis results and chunks
        """
        try:
            # Analyze document
            analysis_result = await self.analyze_document(file_content, file_name)

            # Create chunks from extracted text
            chunks = self.create_chunks(
                text=analysis_result["text"],
                metadata={
                    "file_name": file_name,
                    "source": "document_analysis"
                }
            )

            return {
                "analysis": analysis_result,
                "chunks": chunks,
                "total_chunks": len(chunks)
            }

        except Exception as e:
            logger.error(f"Error in analyze_and_chunk: {str(e)}")
            raise


# Dependency injection
def get_upstage_service() -> UpstageService:
    """Get UpstageService instance"""
    return UpstageService()
